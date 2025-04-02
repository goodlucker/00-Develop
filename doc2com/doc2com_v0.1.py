# !/usr/bin/env python3
# -*- coding: utf-8 -*-
# Author: Rick.Xu

import os
import json
import webbrowser
import pypandoc
from docx import Document
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment, Font
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

CONFIG_FILE = "config.json"

# Namespace dictionary
NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def estimate_page_and_headings(doc):
    """Estimate page number and extract headings for each paragraph."""
    estimated_pages = []
    headings = []
    current_page = 1

    for para in doc.paragraphs:
        # Check if this paragraph is a heading
        if para.style.name.startswith('Heading'):
            headings.append({"text": para.text, "index": doc.paragraphs.index(para), "page": current_page})

        estimated_pages.append({"para": para, "page": current_page})

        # Estimate new page if the paragraph contains a page break
        for run in para.runs:
            if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                current_page += 1

    return estimated_pages, headings

def extract_comments_and_headings_from_docx(docx_path):
    """Extracts comments and headings from a DOCX file."""
    comments = []

    try:
        doc = Document(docx_path)
        rels = None
        for rel in doc.part.rels.values():
            if rel.reltype == RT.COMMENTS:
                rels = rel.target_part
                break

        if not rels:
            return comments  # No comments found

        # Extract comments
        comments_xml = etree.fromstring(rels.blob)
        comment_map = {}
        for comment in comments_xml.findall('.//w:comment', namespaces=NAMESPACE):
            comment_id = comment.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
            author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
            text = ''.join([c.text for c in comment.findall('.//w:t', namespaces=NAMESPACE)])
            comment_map[comment_id] = {"Main Comment": text, "Main Author": author, "Replies": []}

        # Extract replies
        for reply in comments_xml.findall('.//w:comment', namespaces=NAMESPACE):
            parent_id = reply.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}parent")
            if parent_id in comment_map:
                author = reply.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                text = ''.join([c.text for c in reply.findall('.//w:t', namespaces=NAMESPACE)])
                comment_map[parent_id]["Replies"].append({"Reply Author": author, "Reply Text": text})

        estimated_pages, headings = estimate_page_and_headings(doc)

        # Collect comments with page and heading information
        for comment_id, comment in comment_map.items():
            nearest_heading = "No heading"
            comment_page = 1

            # Find the paragraph associated with the comment
            for para_info in estimated_pages:
                para = para_info["para"]
                annotations = para._element.findall('.//w:commentRangeStart[@w:id="{}"]'.format(comment_id), namespaces=NAMESPACE)
                if annotations:
                    comment_page = para_info["page"]
                    for heading in sorted(headings, key=lambda h: h["index"]):
                        if doc.paragraphs.index(para) > heading["index"]:
                            nearest_heading = heading["text"]
                        else:
                            break
                    break

            comments.append({
                "Comment Index": f"Comment {len(comments) + 1}",
                "Annotation Type": "Comment",
                "Main Comment": comment["Main Comment"],
                "Main Author": comment["Main Author"],
                "Replies": ", ".join([f'{r["Reply Author"]}: {r["Reply Text"]}' for r in comment["Replies"]]),
                "Page Number": comment_page,
                "Heading": nearest_heading,
                "File Path": docx_path
            })

    except Exception as e:
        messagebox.showerror("Error", f"Failed to extract annotations and headings from {docx_path}: {str(e)}")
    
    return comments

def convert_rtf_to_docx(rtf_path, docx_path):
    """Convert RTF to DOCX using pypandoc."""
    try:
        pypandoc.convert_file(rtf_path, 'docx', outputfile=docx_path)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to convert RTF to DOCX: {str(e)}")
        raise

def generate_html_link(file_path, comment_idx, html_path):
    """Generate an HTML file that opens the DOCX/RTF file."""
    html_content = f"""
    <html>
    <head>
        <script type="text/javascript">
            function openDoc() {{
                window.open('{file_path}', '_blank');
            }}
        </script>
    </head>
    <body onload="openDoc()"></body>
    </html>
    """
    try:
        with open(html_path, 'w') as file:
            file.write(html_content)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate HTML link: {str(e)}")

def save_annotations_to_excel(annotations, workbook, sheet_name, file_path):
    df = pd.DataFrame(annotations)

    # Create a subdirectory for HTML files if it doesn't exist
    html_dir = os.path.join(os.path.dirname(file_path), "html_links")
    if not os.path.exists(html_dir):
        try:
            os.makedirs(html_dir)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create directory for HTML links: {str(e)}")
            return

    # Get base file name for HTML link naming
    base_file_name = os.path.splitext(os.path.basename(file_path))[0]

    # Create a new sheet
    try:
        worksheet = workbook.create_sheet(title=sheet_name[:31])
    except Exception as e:
        messagebox.showerror("Error", f"Failed to create worksheet: {str(e)}")
        return
    
    for index, row in df.iterrows():
        comment_idx = index + 1  # Comment index
        file_path_formatted = row["File Path"].replace("\\", "/")  # Replace backslashes with forward slashes for URL
        html_file_name = f"{base_file_name}_link_{comment_idx}.html"
        html_file_path = os.path.join(html_dir, html_file_name)

        # Generate HTML file with link to the specific comment
        generate_html_link(file_path_formatted, comment_idx, html_file_path)

        # Create hyperlink to the HTML file
        df.at[index, "Comment Index"] = f'=HYPERLINK("{html_file_path}", "Comment {comment_idx}")'

    # Drop the file path column as it is not needed in the final output
    df.drop(columns=["File Path"], inplace=True)

    # Add data to the specified sheet
    try:
        for row in dataframe_to_rows(df, index=False, header=True):
            worksheet.append(row)
        
        # Apply styles and formats
        for row in worksheet.iter_rows():
            for cell in row:
                # Apply wrap text
                cell.alignment = Alignment(wrap_text=True)
                # Style hyperlinks
                if cell.hyperlink:
                    cell.font = Font(color="0000FF", underline="single")
        
        # Set auto-filter and column widths
        worksheet.auto_filter.ref = worksheet.dimensions
        column_widths = {
            "A": 15,  # Page (Heading)
            "B": 15,  # Comment Index
            "C": 15,  # Annotation Type
            "D": 50,  # Main Comment
            "E": 20,  # Main Author
            "F": 50,  # Replies
            "G": 15,  # Page Number
            "H": 50,  # Heading
        }
        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width
    except Exception as e:
        messagebox.showerror("Error", f"Failed to write data to worksheet: {str(e)}")

def browse_files(filetype):
    try:
        if filetype == "docx":
            files = filedialog.askopenfilenames(filetypes=[("DOCX files", "*.docx")])
        else:
            files = filedialog.askopenfilenames(filetypes=[("RTF files", "*.rtf")])
        files_var.set(";".join(files))
    except Exception as e:
        messagebox.showerror("Error", f"Failed to browse files: {str(e)}")

def browse_output_excel():
    try:
        file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        output_file.set(file)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to browse output file: {str(e)}")

def process_files():
    file_paths = files_var.get().split(";")
    output_path = output_file.get()

    if not file_paths or not output_path:
        messagebox.showwarning("Input Error", "Please select both DOCX/RTF files and output Excel file.")
        return

    try:
        workbook = Workbook()
        for file_path in file_paths:
            try:
                if file_path.endswith('.rtf'):
                    docx_path = file_path.replace('.rtf', '.docx')
                    convert_rtf_to_docx(file_path, docx_path)
                    annotations = extract_comments_and_headings_from_docx(docx_path)
                    os.remove(docx_path)  # Clean up the temp DOCX files after processing
                else:
                    annotations = extract_comments_and_headings_from_docx(file_path)
                file_name = os.path.splitext(os.path.basename(file_path))[0]
                sheet_name = file_name[:31]  # Ensure sheet name is not too long (max 31 characters)
                save_annotations_to_excel(annotations, workbook, sheet_name, file_path)
            except Exception as e:
                messagebox.showerror("Processing Error", f"An error occurred while processing {file_path}: {str(e)}")
        # Remove default sheet created by openpyxl
        if "Sheet" in workbook.sheetnames:
            std = workbook["Sheet"]
            workbook.remove(std)

        try:
            workbook.save(output_path)
            messagebox.showinfo("Success", f"Annotations have been saved to {output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save workbook: {str(e)}")
    except Exception as e:
        messagebox.showerror("Processing Error", f"An error occurred: {str(e)}")

def save_config():
    config = {
        "files": files_var.get(),
        "output_file": output_file.get()
    }
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=4)
        messagebox.showinfo("Success", "Configuration has been saved.")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save configuration: {str(e)}")

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                config = json.load(f)
                files_var.set(config.get("files", ""))
                output_file.set(config.get("output_file", ""))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load configuration: {str(e)}")

def send_feedback():
    try:
        webbrowser.open("mailto:your_email@example.com?subject=Feedback")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open email client: {str(e)}")

# Create the main window
root = tk.Tk()
root.title("DOCX/RTF Annotation Extractor")

# Variables to store file paths
files_var = tk.StringVar()
output_file = tk.StringVar()

# Load previous configuration if it exists
load_config()

# Create and layout the UI elements
tk.Label(root, text="Select DOCX/RTF Files:").grid(row=0, column=0, padx=10, pady=10, sticky="e")
tk.Entry(root, textvariable=files_var, width=50).grid(row=0, column=1, padx=10, pady=10)

tk.Button(root, text="Browse DOCX", command=lambda: browse_files("docx")).grid(row=0, column=2, padx=10, pady=10)
tk.Button(root, text="Browse RTF", command=lambda: browse_files("rtf")).grid(row=0, column=3, padx=10, pady=10)

tk.Label(root, text="Select Output Excel File:").grid(row=1, column=0, padx=10, pady=10, sticky="e")
tk.Entry(root, textvariable=output_file, width=50).grid(row=1, column=1, padx=10, pady=10)
tk.Button(root, text="Browse", command=browse_output_excel).grid(row=1, column=2, padx=10, pady=10)

tk.Button(root, text="Extract Annotations", command=process_files).grid(row=2, columnspan=3, pady=10)
tk.Button(root, text="Save Configuration", command=save_config).grid(row=3, columnspan=3, pady=10)

feedback_button = tk.Button(root, text="Feedback", command=send_feedback)
feedback_button.place(relx=1.0, rely=1.0, anchor="se", x=-10, y=-10)

# Run the main loop
root.mainloop()